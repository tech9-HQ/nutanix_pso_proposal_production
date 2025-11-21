# app/services/professional_styling.py
"""
Professional styling enhancements for proposals
Adds magazine-quality visual elements
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict
import datetime

# ============================================================================
# COLOR SCHEME
# ============================================================================

class Colors:
    """Professional color palette - Tech9Labs branding"""
    PRIMARY_DARK = RGBColor(0, 51, 102)      # Dark blue
    PRIMARY_MEDIUM = RGBColor(0, 102, 204)   # Medium blue
    PRIMARY_LIGHT = RGBColor(230, 242, 255)  # Light blue
    
    ACCENT_GREEN = RGBColor(0, 170, 0)       # Success/positive
    ACCENT_ORANGE = RGBColor(255, 153, 0)    # Warning
    ACCENT_RED = RGBColor(192, 0, 0)         # Critical
    
    TEXT_DARK = RGBColor(68, 68, 68)         # Body text
    TEXT_LIGHT = RGBColor(128, 128, 128)     # Captions
    
    BG_LIGHT_GRAY = 'F2F2F2'                 # Alternating rows
    BG_LIGHT_BLUE = 'E6F2FF'                 # Info boxes
    BG_WHITE = 'FFFFFF'

# ============================================================================
# ENHANCED COVER PAGE
# ============================================================================

def add_magazine_cover_page(
    doc: Document, 
    customer_name: str, 
    industry: str,
    logo_path: str = "app/assets/tech9labs_logo.png"
) -> None:
    """Create magazine-quality cover page"""
    
    # Top colored bar with title
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    
    # Dark blue background
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '003366')
    header_cell._element.get_or_add_tcPr().append(shading)
    
    # White text in header
    header_para = header_cell.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("PROFESSIONAL SERVICES PROPOSAL")
    header_run.bold = True
    header_run.font.size = Pt(22)
    header_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Padding for header
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'bottom']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), '200')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    header_cell._element.get_or_add_tcPr().append(tcMar)
    
    # Spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Large customer name
    customer_para = doc.add_paragraph()
    customer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    customer_run = customer_para.add_run(customer_name)
    customer_run.bold = True
    customer_run.font.size = Pt(36)
    customer_run.font.color.rgb = Colors.PRIMARY_DARK
    
    # Subtitle
    doc.add_paragraph()
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Nutanix Infrastructure Transformation")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = Colors.TEXT_DARK
    
    if industry:
        industry_para = doc.add_paragraph()
        industry_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        industry_run = industry_para.add_run(f"{industry} Industry")
        industry_run.font.size = Pt(14)
        industry_run.italic = True
        industry_run.font.color.rgb = Colors.TEXT_LIGHT
    
    # Horizontal line
    doc.add_paragraph()
    add_horizontal_line(doc, color=Colors.PRIMARY_MEDIUM, thickness=2)
    doc.add_paragraph()
    
    # Centered logo
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Inches(3.0))
    except Exception:
        logo_text = logo_para.add_run("Integrated Tech9 Labs")
        logo_text.font.size = Pt(24)
        logo_text.bold = True
        logo_text.font.color.rgb = Colors.PRIMARY_DARK
    
    # Bottom details
    for _ in range(3):
        doc.add_paragraph()
    
    # Details table
    details_table = doc.add_table(rows=3, cols=2)
    details_table.style = 'Light Grid Accent 1'
    
    today = datetime.date.today().strftime("%B %d, %Y")
    details = [
        ("Industry Sector", industry or "Technology Services"),
        ("Proposal Date", today),
        ("Prepared By", "Integrated Tech9 Labs Pvt. Ltd.")
    ]
    
    for i, (label, value) in enumerate(details):
        label_cell = details_table.rows[i].cells[0]
        value_cell = details_table.rows[i].cells[1]
        
        label_cell.text = label
        for para in label_cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = Colors.PRIMARY_DARK
        
        value_cell.text = value
        for para in value_cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = Colors.TEXT_DARK
    
    # Confidentiality notice
    doc.add_paragraph()
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run("CONFIDENTIAL — For Client Review Only")
    conf_run.bold = True
    conf_run.font.size = Pt(10)
    conf_run.font.color.rgb = Colors.ACCENT_RED
    
    doc.add_page_break()


def _color_to_hex(color) -> str:
    """
    Convert an RGBColor or hex string into 'RRGGBB' for Word.
    """
    # If already hex string
    if isinstance(color, str):
        return color.lstrip("#").upper()

    # If docx RGBColor instance
    if isinstance(color, RGBColor):
        # RGBColor → string "0xRRGGBB"
        s = str(color)  # e.g. "0x0033CC"
        if s.startswith("0x"):
            s = s[2:]
        return s.upper()

    # fallback
    return "003366"


def add_horizontal_line(doc: Document, color=Colors.PRIMARY_DARK, thickness: int = 1) -> None:
    """
    Add horizontal line separator with correct color handling.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness * 8))
    bottom.set(qn('w:space'), '1')

    # FIXED — use safe hex converter
    bottom.set(qn('w:color'), _color_to_hex(color))

    pBdr.append(bottom)
    pPr.append(pBdr)



def add_section_divider(doc: Document, title: str) -> None:
    """Add colored section divider bar"""
    
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '0066CC')
    cell._element.get_or_add_tcPr().append(shading)
    
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'bottom']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), '100')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    cell._element.get_or_add_tcPr().append(tcMar)
    
    doc.add_paragraph()


def add_callout_box(doc: Document, title: str, content: str, box_type: str = "info") -> None:
    """Add colored call-out box"""
    
    schemes = {
        "info": {"bg": "E6F2FF", "border": "0066CC"},
        "warning": {"bg": "FFF4E6", "border": "FF9900"},
        "success": {"bg": "E6F9E6", "border": "00AA00"},
        "tip": {"bg": "F0E6FF", "border": "9933CC"}
    }
    
    scheme = schemes.get(box_type, schemes["info"])
    
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), scheme["bg"])
    cell._element.get_or_add_tcPr().append(shading)
    
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), scheme["border"])
    tcBorders.append(left)
    cell._element.get_or_add_tcPr().append(tcBorders)
    
    title_para = cell.paragraphs[0]
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(12)
    
    content_para = cell.add_paragraph()
    content_para.add_run(content)
    
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), '150')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    cell._element.get_or_add_tcPr().append(tcMar)


# ============================================================================
# ENHANCED TABLES
# ============================================================================

def create_professional_table(
    doc: Document, 
    headers: List[str], 
    data: List[List[str]],
    highlight_totals: bool = False
) -> object:
    """Create professionally styled table with alternating rows"""
    
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row - dark blue with white text
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003366')
        cell._element.get_or_add_tcPr().append(shading)
        
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows with alternating colors
    for row_idx, row_data in enumerate(data):
        row = table.add_row()
        
        # Alternate row colors
        if row_idx % 2 == 1:
            for cell in row.cells:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), Colors.BG_LIGHT_GRAY)
                cell._element.get_or_add_tcPr().append(shading)
        
        is_total_row = highlight_totals and row_idx == len(data) - 1
        
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            
            run = para.add_run(str(value))
            
            if is_total_row:
                run.bold = True
                run.font.color.rgb = Colors.PRIMARY_DARK
    
    return table


def add_key_metrics_grid(doc: Document, metrics: List[Dict[str, str]]) -> None:
    """Add visual grid of key metrics (2-column layout)"""
    
    table = doc.add_table(rows=(len(metrics) + 1) // 2, cols=2)
    
    for idx, metric in enumerate(metrics):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = table.rows[row_idx].cells[col_idx]
        
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), Colors.BG_LIGHT_GRAY)
        cell._element.get_or_add_tcPr().append(shading)
        
        # Large value
        value_para = cell.paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.add_run(metric['value'])
        value_run.bold = True
        value_run.font.size = Pt(24)
        value_run.font.color.rgb = Colors.PRIMARY_MEDIUM
        
        # Small label
        label_para = cell.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.add_run(metric['label'])
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = Colors.TEXT_LIGHT
        
        tcMar = OxmlElement('w:tcMar')
        for margin in ['top', 'bottom']:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), '150')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        cell._element.get_or_add_tcPr().append(tcMar)