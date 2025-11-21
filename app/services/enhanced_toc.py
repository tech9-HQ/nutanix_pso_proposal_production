# app/services/enhanced_toc.py
"""
Enhanced Table of Contents with Professional Styling
Integrates magazine-quality visual elements
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Tuple, Dict
import io
import logging

# Import professional styling
from app.services.professional_styling import (
    add_magazine_cover_page,
    add_section_divider,
    Colors
)

log = logging.getLogger("enhanced_toc_builder")

# ============================================================================
# SMART SECTION GROUPING
# ============================================================================

DETAILED_PROPOSAL_STRUCTURE: List[Tuple[str, str, List[Tuple[str, str]]]] = [
    ("1", "Executive Overview", [
        ("cover_page", "Cover Page"),
        ("executive_summary", "Executive Summary"),
        ("about_tech9labs", "About Integrated Tech9 Labs"),
    ]),
    ("2", "Customer Context & Requirements", [
        ("customer_background_and_business_drivers", "Customer Background & Business Drivers"),
        ("current_infrastructure_assessment", "Current Infrastructure Assessment"),
    ]),
    ("3", "Technical Solution Architecture", [
        ("target_state_architecture", "Target State Architecture"),
        ("migration_strategy_and_approach", "Migration Strategy & Approach"),
        ("parallel_build_approach", "Parallel Build Methodology"),
        ("migration_waves", "Wave-Based Migration Plan"),
        ("rollback_strategy", "Rollback & Contingency Strategy"),
        ("validation_strategy", "Validation & Testing Strategy"),
    ]),
    ("4", "Scope of Work & Deliverables", [
        ("scope_of_work_in_scope", "In-Scope Activities"),
        ("scope_of_work_out_of_scope", "Out-of-Scope Items"),
        ("detailed_wbs", "Work Breakdown Structure (WBS)"),
        ("tools_and_technologies", "Tools & Technologies"),
        ("deployment_and_configuration_details", "Deployment & Configuration Details"),
        ("testing_validation_and_acceptance", "Testing, Validation & Acceptance Criteria"),
        ("project_deliverables", "Project Deliverables"),
    ]),
    ("5", "Project Governance & Management", [
        ("project_governance", "Governance Framework"),
        ("raci_matrix", "RACI Matrix"),
        ("communication_plan", "Communication Plan"),
        ("escalation_matrix", "Escalation Procedures"),
    ]),
    ("6", "Risk Management & Dependencies", [
        ("assumptions_and_dependencies", "Assumptions & Dependencies"),
        ("risks_and_mitigation", "Risk Register & Mitigation Strategies"),
    ]),
    ("7", "Commercial & Timeline", [
        ("commercial_boq_expanded", "Commercial Bill of Quantities"),
        ("project_timeline", "Project Timeline & Milestones"),
    ]),
    ("8", "Annexures & References", [
        ("annexures", "Annexures & Supporting Documentation"),
    ]),
]


def add_enhanced_table_of_contents(doc: Document) -> None:
    """Add professional categorized TOC with styling"""
    
    # Title
    toc_title = doc.add_heading("Table of Contents", level=0)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = Colors.PRIMARY_DARK
        run.bold = True
    
    doc.add_paragraph()
    
    # Instructional note
    note = doc.add_paragraph()
    note_run = note.add_run(
        "This proposal is organized into 8 major sections. "
        "Click any section title to navigate directly in Microsoft Word."
    )
    note_run.font.size = Pt(10)
    note_run.italic = True
    note_run.font.color.rgb = Colors.TEXT_LIGHT
    
    doc.add_paragraph()
    
    # Manual TOC with styling
    for section_num, category_name, subsections in DETAILED_PROPOSAL_STRUCTURE:
        # Category heading
        category_para = doc.add_paragraph()
        category_para.paragraph_format.space_before = Pt(12)
        category_para.paragraph_format.space_after = Pt(6)
        
        category_run = category_para.add_run(f"{section_num}. {category_name}")
        category_run.bold = True
        category_run.font.size = Pt(13)
        category_run.font.color.rgb = Colors.PRIMARY_DARK
        
        # Subsections
        for idx, (section_key, section_title) in enumerate(subsections, start=1):
            subsection_para = doc.add_paragraph()
            subsection_para.paragraph_format.left_indent = Inches(0.3)
            subsection_para.paragraph_format.space_before = Pt(2)
            subsection_para.paragraph_format.space_after = Pt(2)
            
            num_run = subsection_para.add_run(f"{section_num}.{idx} ")
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = Colors.TEXT_DARK
            
            title_run = subsection_para.add_run(section_title)
            title_run.font.size = Pt(11)
            
            leader_run = subsection_para.add_run(" " + "." * 80)
            leader_run.font.color.rgb = RGBColor(200, 200, 200)
            leader_run.font.size = Pt(9)
            
            page_run = subsection_para.add_run(" XX")
            page_run.font.size = Pt(11)
            page_run.font.color.rgb = Colors.TEXT_DARK
    
    doc.add_paragraph()
    
    # Word TOC field
    info_para = doc.add_paragraph()
    info_run = info_para.add_run(
        "Note: Right-click the table below and select 'Update Field' "
        "to populate accurate page numbers."
    )
    info_run.font.size = Pt(9)
    info_run.italic = True
    info_run.font.color.rgb = Colors.TEXT_LIGHT
    
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run()
    
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    toc_run._r.append(fld_char_begin)
    
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    toc_run._r.append(instr_text)
    
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    toc_run._r.append(fld_char_separate)
    
    toc_run._r.append(OxmlElement("w:t"))
    toc_run._r[-1].text = "[Automatic TOC - Update this field in Word]"
    
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    toc_run._r.append(fld_char_end)
    
    doc.add_page_break()


def get_section_hierarchy() -> Dict[str, Tuple[int, int]]:
    """Return section hierarchy mapping"""
    hierarchy = {}
    for cat_idx, (section_num, category_name, subsections) in enumerate(DETAILED_PROPOSAL_STRUCTURE, start=1):
        for sub_idx, (section_key, section_title) in enumerate(subsections, start=1):
            hierarchy[section_key] = (cat_idx, sub_idx)
    return hierarchy


def add_section_with_smart_heading(
    doc: Document,
    section_key: str,
    title: str,
    content: str,
    hierarchy: Dict[str, Tuple[int, int]],
    add_page_break: bool = True,
) -> None:
    """Add section with smart heading levels"""
    
    if section_key in hierarchy:
        cat_num, sub_num = hierarchy[section_key]
        is_category_start = sub_num == 1
        
        if is_category_start:
            category_title = None
            for section_num, cat_name, subsections in DETAILED_PROPOSAL_STRUCTURE:
                if subsections[0][0] == section_key:
                    category_title = f"{section_num}. {cat_name}"
                    break
            
            if category_title:
                # Add visual section divider
                add_section_divider(doc, cat_name)
        
        section_heading = doc.add_heading(title, level=2)
        section_heading.paragraph_format.space_before = Pt(12)
        section_heading.paragraph_format.space_after = Pt(6)
    else:
        doc.add_heading(title, level=2)
    
    if content and content.strip():
        import re
        from app.services.docx_builder import add_block_as_paragraphs
        
        lines = [l.rstrip() for l in content.splitlines()]
        if any(line.startswith("|") for line in lines):
            add_block_as_paragraphs(doc, content)
        else:
            blocks = re.split(r"\n\s*\n", content)
            for block in blocks:
                block = block.strip()
                if block:
                    add_block_as_paragraphs(doc, block)
    
    if add_page_break:
        doc.add_page_break()


# ============================================================================
# MAIN BUILDER FUNCTION
# ============================================================================

def build_detailed_proposal_with_enhanced_toc(sections: Dict[str, str]) -> bytes:
    """Build detailed proposal with professional styling"""
    
    log.info("="*80)
    log.info("BUILDING PROFESSIONAL DETAILED PROPOSAL")
    log.info("="*80)
    
    # Clean metadata
    cleaned_sections = {}
    metadata_keys = {
        "_metadata", "_proposal_type", "metadata", "proposal_type", 
        "qa_report", "debug_info", "error", "parse_error", "raw_output", 
        "exception", "final_answer"
    }
    
    for key, value in sections.items():
        if key not in metadata_keys:
            cleaned_sections[key] = value
        else:
            log.info(f"Filtered: {key}")
    
    sections = cleaned_sections
    log.info(f"Cleaned: {len(sections)} sections")
    
    doc = Document()
    
    from app.services.docx_builder import (
        apply_base_styles,
        add_branding,
        fetch_exchange_rate,
        add_commercial_boq_section,
    )
    
    apply_base_styles(doc)
    add_branding(doc)
    
    # Extract customer info
    cover_text = sections.get("cover_page", "")
    customer_name = "Customer"
    industry = "Technology Services"
    
    if "BirlaSoft" in str(sections.values()):
        customer_name = "BirlaSoft Limited"
        industry = "ITES"
    elif "customer" in cover_text.lower():
        import re
        match = re.search(r"customer[:\s]+([^\n]+)", cover_text, re.IGNORECASE)
        if match:
            customer_name = match.group(1).strip()
    
    # Professional cover page
    log.info(f"Adding professional cover for {customer_name}")
    add_magazine_cover_page(doc, customer_name, industry)
    
    # Enhanced TOC
    log.info("Adding enhanced TOC...")
    try:
        add_enhanced_table_of_contents(doc)
        log.info("✓ TOC added")
    except Exception as e:
        log.error(f"✗ TOC failed: {e}")
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("Update field in Word for page numbers")
        doc.add_page_break()
    
    # Content sections
    hierarchy = get_section_hierarchy()
    section_order = []
    for _, _, subsections in DETAILED_PROPOSAL_STRUCTURE:
        for section_key, _ in subsections:
            if section_key != "cover_page":
                section_order.append(section_key)
    
    log.info(f"Adding {len(section_order)} sections...")
    
    exchange_rate = fetch_exchange_rate()
    
    for idx, section_key in enumerate(section_order):
        content = sections.get(section_key, "")
        if not content or not content.strip():
            log.warning(f"⚠ Missing: {section_key}")
            continue
        
        title = None
        for _, _, subsections in DETAILED_PROPOSAL_STRUCTURE:
            for key, section_title in subsections:
                if key == section_key:
                    title = section_title
                    break
            if title:
                break
        
        if not title:
            title = section_key.replace("_", " ").title()
        
        is_last = (idx == len(section_order) - 1)
        
        if section_key == "commercial_boq_expanded":
            cat_num, sub_num = hierarchy.get(section_key, (7, 1))
            if sub_num == 1:
                add_section_divider(doc, "Commercial & Timeline")
            
            log.info(f"Adding BOQ: {len(content)} chars")
            add_commercial_boq_section(
                doc, title, content,
                exchange_rate=exchange_rate,
                heading_level=2,
            )
        else:
            add_section_with_smart_heading(
                doc, section_key, title, content,
                hierarchy, add_page_break=not is_last,
            )
        
        log.info(f"✓ Added {idx+1}/{len(section_order)}: {title}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    log.info("="*80)
    log.info("✓ PROFESSIONAL PROPOSAL COMPLETE")
    log.info("="*80)
    
    return buffer.read()