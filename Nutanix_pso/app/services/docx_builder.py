"""
Professional DOCX Builder for Tech9Labs Proposals
Generates detailed and short-form proposals with proper formatting,
cost calculations in USD/INR, and comprehensive styling.
"""

from __future__ import annotations

import io
import os
import re
import datetime
import logging
from decimal import Decimal, ROUND_HALF_UP
from typing import Dict, List, Any, Tuple, Optional
from dataclasses import dataclass

import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================================
# CONFIGURATION & CONSTANTS
# ============================================================================

log = logging.getLogger("docx_builder")

LOGO_PATH = "app/assets/tech9labs_logo.png"

# Base default
DEFAULT_RATE_USD = Decimal("400")  # USD per man-day (generic default)

# Category-based average rates (USD / man-day)
CATEGORY_RATE_USD: Dict[str, Decimal] = {
    "Migration Services": Decimal("400"),
    "Assessment Services": Decimal("600"),
    "Development Services": Decimal("600"),
    "AI & Advanced Analytics": Decimal("600"),
    "Deployment Services": Decimal("400"),
}

# Any workshops
WORKSHOP_RATE_USD = Decimal("600")

FALLBACK_USD_INR_RATE = Decimal("87.95")
FX_API_TIMEOUT = 5  # seconds

# Section title mappings for detailed proposals
SECTION_TITLES: Dict[str, str] = {
    "cover_page": "Cover Page",
    "executive_summary": "Executive Summary",
    "about_tech9labs": "About Tech9Labs",
    "customer_background_and_business_drivers": "Customer Background & Business Drivers",
    "current_infrastructure_assessment": "Current Infrastructure Assessment",
    "target_state_architecture": "Target State Architecture",
    "migration_strategy_and_approach": "Migration Strategy & Approach (Parallel Build)",
    "parallel_build_approach": "Parallel Build Approach",
    "migration_waves": "Migration Waves",
    "rollback_strategy": "Rollback Strategy",
    "validation_strategy": "Validation Strategy",
    "scope_of_work_in_scope": "Scope of Work – In Scope",
    "scope_of_work_out_of_scope": "Scope of Work – Out of Scope",
    "detailed_wbs": "Work Breakdown Structure (WBS)",
    "tools_and_technologies": "Tools & Technologies Used",
    "deployment_and_configuration_details": "Deployment & Configuration Details",
    "testing_validation_and_acceptance": "Testing, Validation & Acceptance Criteria",
    "project_deliverables": "Project Deliverables",
    "project_governance": "Project Governance",
    "raci_matrix": "RACI Matrix",
    "communication_plan": "Communication Plan",
    "escalation_matrix": "Escalation Matrix",
    "assumptions_and_dependencies": "Assumptions & Dependencies",
    "risks_and_mitigation": "Risks & Mitigation",
    "commercial_boq_expanded": "Commercial Bill of Quantities",
    "project_timeline": "Project Timeline (Gantt Style)",
    "annexures": "Annexures",
}

# Grouped headings for ToC (Heading 1)
TOC_GROUPS: List[Tuple[str, List[str]]] = [
    (
        "1. Assessment & Design",
        [
            "executive_summary",
            "about_tech9labs",
            "customer_background_and_business_drivers",
            "current_infrastructure_assessment",
            "target_state_architecture",
        ],
    ),
    (
        "2. Migration & Cutover",
        [
            "migration_strategy_and_approach",
            "parallel_build_approach",
            "migration_waves",
            "rollback_strategy",
            "validation_strategy",
        ],
    ),
    (
        "3. Scope & Deliverables",
        [
            "scope_of_work_in_scope",
            "scope_of_work_out_of_scope",
            "detailed_wbs",
            "project_deliverables",
        ],
    ),
    (
        "4. Governance, Risks & Communication",
        [
            "project_governance",
            "raci_matrix",
            "communication_plan",
            "escalation_matrix",
            "assumptions_and_dependencies",
            "risks_and_mitigation",
        ],
    ),
    (
        "5. Commercials & Timeline",
        [
            "commercial_boq_expanded",
            "project_timeline",
        ],
    ),
    (
        "6. Annexures",
        [
            "annexures",
        ],
    ),
]

def _find_group_for_key(section_key: str) -> Optional[str]:
    for group_title, keys in TOC_GROUPS:
        if section_key in keys:
            return group_title
    return None

# Default Terms & Conditions
TERMS_AND_CONDITIONS: List[str] = [
    "All services are delivered on a best-effort basis in accordance with the agreed Statement of Work.",
    "Pricing is exclusive of applicable taxes (GST/VAT) and government duties.",
    "Payment terms: 50% advance, 50% upon project completion, unless otherwise agreed.",
    "Travel and accommodation expenses, if required, will be charged on actuals with prior approval.",
    "Client shall provide necessary access, credentials, and resources for successful project execution.",
    "Any scope changes will be managed through a formal Change Request process with revised timelines and costs.",
]

# Regex patterns for markdown parsing
BOLD_PATTERN = re.compile(r"(\*\*.*?\*\*)")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)$")


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class CostItem:
    """Represents a single cost line item in the BOQ"""
    description: str
    man_days: int
    rate_usd: Decimal
    total_usd: Decimal
    total_inr: Decimal


@dataclass
class ExchangeRate:
    """Exchange rate information"""
    rate: Decimal
    source: str
    timestamp: datetime.datetime


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def sanitize_filename(name: Optional[str]) -> str:
    """
    Sanitize filename by removing special characters.
    """
    if not name:
        return "proposal"
    
    base = re.sub(r"[^A-Za-z0-9._\-]+", "_", name.strip())
    return base or "proposal"


def safe_save_doc(doc: Document, filename: str) -> str:
    """
    Save document with automatic versioning if file exists.
    """
    base_dir = os.getenv("OUTPUT_DIR", "generated_docs")
    os.makedirs(base_dir, exist_ok=True)

    name, ext = os.path.splitext(filename or "proposal.docx")
    ext = ext or ".docx"

    candidate = os.path.join(base_dir, name + ext)
    counter = 1
    
    while os.path.exists(candidate):
        candidate = os.path.join(base_dir, f"{name}_{counter}{ext}")
        counter += 1

    doc.save(candidate)
    log.info(f"Document saved successfully: {candidate}")
    return candidate


def fetch_exchange_rate() -> ExchangeRate:
    """
    Fetch current USD to INR exchange rate from API.
    Falls back to configured default if API fails.

    Business rule:
    - Always add 1.00 INR to the live/fallback rate before using it.
    """
    def _apply_buffer(raw: Decimal) -> Decimal:
        return (raw + Decimal("1")).quantize(Decimal("0.01"), ROUND_HALF_UP)

    try:
        response = requests.get(
            "https://api.exchangerate.host/latest",
            params={"base": "USD", "symbols": "INR"},
            timeout=FX_API_TIMEOUT,
        )
        response.raise_for_status()
        data = response.json()
        
        if data and data.get("rates") and data["rates"].get("INR"):
            raw_rate = Decimal(str(data["rates"]["INR"]))
            effective_rate = _apply_buffer(raw_rate)
            log.info(
                f"Fetched live exchange rate: 1 USD = {raw_rate} INR "
                f"(effective billing rate: {effective_rate} INR)"
            )
            return ExchangeRate(
                rate=effective_rate,
                source="exchangerate.host API +1 INR buffer",
                timestamp=datetime.datetime.now(),
            )
    except Exception as e:
        log.warning(f"Exchange rate fetch failed: {e}")

    # Fallback
    raw_fallback = Decimal(os.getenv("FALLBACK_USD_INR", str(FALLBACK_USD_INR_RATE)))
    effective_fallback = _apply_buffer(raw_fallback)
    log.info(
        f"Using fallback exchange rate: 1 USD = {raw_fallback} INR "
        f"(effective billing rate: {effective_fallback} INR)"
    )
    return ExchangeRate(
        rate=effective_fallback,
        source="Fallback configuration +1 INR buffer",
        timestamp=datetime.datetime.now(),
    )


def format_currency_usd(amount: Decimal) -> str:
    return f"${amount:,.2f}"


def format_currency_inr(amount: Decimal) -> str:
    return f"₹{amount:,.2f}"


# ============================================================================
# DOCUMENT STYLING
# ============================================================================

def apply_base_styles(doc: Document) -> None:
    """
    Apply consistent professional base styles to document.
    """
    try:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.15
    except Exception as e:
        log.error(f"Failed to apply Normal style: {e}")

    heading_configs = {
        "Heading 1": {"size": Pt(16), "bold": True, "color": RGBColor(0, 51, 102)},
        "Heading 2": {"size": Pt(14), "bold": True, "color": RGBColor(0, 51, 102)},
        "Heading 3": {"size": Pt(12), "bold": True, "color": RGBColor(68, 68, 68)},
        "Heading 4": {"size": Pt(11), "bold": True, "color": RGBColor(68, 68, 68)},
    }
    
    for heading_name, config in heading_configs.items():
        try:
            style = doc.styles[heading_name]
            style.font.name = "Calibri"
            style.font.size = config["size"]
            style.font.bold = config["bold"]
            style.font.color.rgb = config["color"]
            
            pf = style.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.keep_with_next = True
        except Exception as e:
            log.warning(f"Could not style {heading_name}: {e}")


def compact_table(table) -> None:
    """
    Remove unnecessary spacing from table cells.
    """
    try:
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                pf = paragraph.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)
                pf.line_spacing = 1.0


def add_inline_bold_runs(paragraph, text: str) -> None:
    """
    Parse and render **bold** markdown syntax in paragraph.
    """
    parts = BOLD_PATTERN.split(text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run_text = part[2:-2]
            run = paragraph.add_run(run_text)
            run.bold = True
        else:
            paragraph.add_run(part)


# ============================================================================
# BRANDING & HEADERS
# ============================================================================

def add_branding(document: Document) -> None:
    """
    Add Tech9Labs branding to document header and footer.
    """
    section = document.sections[0]

    # HEADER
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    try:
        run = header_para.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.5))
    except Exception as e:
        log.warning(f"Could not add logo: {e}")
        header_para.add_run("Integrated Tech9 Labs Pvt. Ltd.")

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "Confidential — © Integrated Tech9 Labs Pvt. Ltd. | Not for Distribution"
    
    if footer_para.runs:
        run = footer_para.runs[0]
    else:
        run = footer_para.add_run()
    
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)


# ============================================================================
# COVER PAGE FUNCTIONS
# ============================================================================

def parse_cover_page_text(cover_text: str) -> Tuple[str, str, str, List[str]]:
    """
    Extract structured information from cover page text.
    """
    lines = [l.strip() for l in cover_text.splitlines() if l.strip()]
    project_title = "Nutanix Professional Services Proposal"
    customer_name = ""
    prepared_by = "Integrated Tech9Labs Pvt. Ltd."
    contents_items: List[str] = []

    if lines:
        project_title = lines[0]

    for idx, line in enumerate(lines):
        lower = line.lower()
        
        if lower.startswith("project title:"):
            project_title = line.split(":", 1)[1].strip() or project_title
        elif lower.startswith("customer:"):
            customer_name = line.split(":", 1)[1].strip()
        elif lower.startswith("prepared by:"):
            prepared_by = line.split(":", 1)[1].strip() or prepared_by
        elif lower.startswith("contents:"):
            for j in range(idx + 1, len(lines)):
                item = lines[j]
                if item.startswith(("-", "•", "*")):
                    contents_items.append(re.sub(r"^[-*•]\s*", "", item).strip())
                else:
                    break

    normalized = project_title.strip().lower()
    if normalized in {"cover page", "cover"}:
        if customer_name:
            project_title = f"Nutanix Professional Services Proposal for {customer_name}"
        else:
            project_title = "Nutanix Professional Services Proposal"

    return project_title, customer_name, prepared_by, contents_items


def guess_customer_from_sections(sections: Dict[str, str]) -> str:
    """
    Attempt to extract customer name from section content.
    """
    full_text = "\n".join(sections.values())
    
    org_pattern = r"\b([A-Z][A-Za-z0-9&., ]+\b(?:Limited|Ltd\.?|Pvt\.?|Inc\.?|LLC|Technologies|Systems|Corporation|Corp\.?))"
    match = re.search(org_pattern, full_text)
    
    if match:
        return match.group(1).strip()

    match2 = re.search(r"\b([A-Z][A-Za-z0-9]+ (?:Limited|Technologies))\b", full_text)
    if match2:
        return match2.group(1).strip()

    return ""


def add_detailed_cover_page(
    doc: Document,
    cover_text: str,
    sections: Dict[str, str]
) -> Tuple[str, str]:
    """
    Create professional cover page for detailed proposals.
    """
    add_branding(doc)

    project_title, customer_name, prepared_by, contents_items = parse_cover_page_text(cover_text)
    
    if not customer_name:
        guessed = guess_customer_from_sections(sections)
        if guessed:
            customer_name = guessed

    today_str = datetime.date.today().strftime("%d %B %Y")

    heading = doc.add_heading("Nutanix Professional Services Proposal", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    if customer_name:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(f"Prepared for: {customer_name}")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph()  # Spacing

    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    
    labels = ["Project Title", "Customer", "Prepared By", "Date"]
    values = [project_title, customer_name or "N/A", prepared_by, today_str]

    for i in range(4):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]

        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(labels[i] + ":")
        label_run.bold = True
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        value_para = value_cell.paragraphs[0]
        value_para.add_run(values[i])
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    compact_table(table)

    if contents_items:
        doc.add_paragraph()
        doc.add_heading("Document Contents", level=2)
        for item in contents_items:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(item)

    doc.add_page_break()

    return customer_name, project_title


def add_premium_cover_page(
    doc: Document,
    customer_name: str,
    industry: Optional[str],
) -> None:
    """
    Create premium cover page for short proposals.
    """
    add_branding(doc)

    heading = doc.add_heading("Nutanix Professional Services Proposal", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"Prepared for: {customer_name}")
    sub_run.bold = True
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0, 102, 204)

    if industry:
        ind_para = doc.add_paragraph()
        ind_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ind_run = ind_para.add_run(f"Industry: {industry}")
        ind_run.font.size = Pt(12)
        ind_run.italic = True

    doc.add_paragraph()  # Spacing

    today_str = datetime.date.today().strftime("%d %B %Y")
    
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(f"Prepared By: Integrated Tech9 Labs Pvt. Ltd.\nDate: {today_str}")

    doc.add_paragraph()  # Spacing

    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run("CONFIDENTIAL – For Client Review Only")
    conf_run.bold = True
    conf_run.font.size = Pt(10)
    conf_run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()
    log.info(f"Premium cover page added for {customer_name}")


# ============================================================================
# TABLE OF CONTENTS
# ============================================================================

def add_table_of_contents(doc: Document, section_titles: Optional[List[str]] = None) -> None:
    """
    Insert professional Table of Contents.
    The real page numbers come from the Word TOC field (Heading 1–3).
    """
    if section_titles is None:
        section_titles = []

    toc_heading = doc.add_paragraph()
    toc_heading.style = "Heading 1"
    toc_heading.add_run("Table of Contents")

    # Optional manual overview (not linked)
    for idx, title in enumerate(section_titles, start=1):
        para = doc.add_paragraph()
        para.style = "Normal"
        run = para.add_run(f"{idx}. {title}")
        run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    p = doc.add_paragraph()
    r = p.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    r._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    r._r.append(fld_char_separate)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_char_end)

    doc.add_page_break()


# ============================================================================
# MARKDOWN CONTENT RENDERING
# ============================================================================

def looks_like_md_table(lines: List[str]) -> bool:
    if len(lines) < 2:
        return False
    
    if not (lines[0].startswith("|") and lines[0].endswith("|")):
        return False
    
    separator = lines[1].strip()
    if not (separator.startswith("|") and separator.endswith("|")):
        return False
    
    core = separator.replace("|", "").replace(" ", "")
    return bool(core) and all(ch in "-:" for ch in core)


def render_md_table(doc: Document, lines: List[str]) -> None:
    header_line = lines[0].strip()
    data_lines = [l.strip() for l in lines[2:] if l.strip().startswith("|")]

    headers = [h.strip() for h in header_line.strip("|").split("|")]
    num_cols = len(headers)
    
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        for para in header_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True

    for data_line in data_lines:
        parts = [c.strip() for c in data_line.strip("|").split("|")]
        row_cells = table.add_row().cells
        
        for i in range(num_cols):
            row_cells[i].text = parts[i] if i < len(parts) else ""

    compact_table(table)


def add_block_as_paragraphs(doc: Document, block: str) -> None:
    """
    Intelligently render text block with markdown support.
    """
    raw_lines = [l.rstrip() for l in block.splitlines()]
    lines = [l for l in raw_lines if l.strip()]
    
    if not lines:
        return

    # Markdown table
    if looks_like_md_table(lines):
        render_md_table(doc, lines)
        return

    # Bullet list
    if all(l.lstrip().startswith(("-", "*", "•")) for l in lines):
        for line in lines:
            text = re.sub(r"^[-*•]\s*", "", line.lstrip())
            if text:
                para = doc.add_paragraph(style="List Bullet")
                add_inline_bold_runs(para, text)
        return

    # Numbered list
    if all(re.match(r"^\s*\d+\.", l) for l in lines):
        for line in lines:
            text = re.sub(r"^\s*\d+\.\s*", "", line)
            if text:
                para = doc.add_paragraph(style="List Number")
                add_inline_bold_runs(para, text)
        return

    # Markdown headings
    if any(HEADING_PATTERN.match(l) for l in lines):
        for line in lines:
            match = HEADING_PATTERN.match(line)
            if match:
                hashes, text = match.groups()
                text = text.strip()
                if text:
                    level = min(len(hashes) + 1, 4)
                    para = doc.add_paragraph(style=f"Heading {level}")
                    add_inline_bold_runs(para, text)
            else:
                para = doc.add_paragraph()
                add_inline_bold_runs(para, line.strip())
        return

    # Normal paragraph
    joined_text = " ".join(l.strip() for l in lines)
    para = doc.add_paragraph()
    add_inline_bold_runs(para, joined_text)


def add_section(
    doc: Document,
    title: str,
    content: str,
    add_page_break: bool = True,
    heading_level: int = 1,
) -> None:
    """
    Add a complete document section with title and content.
    heading_level controls Heading 1/2 etc. for grouped ToC.
    """
    doc.add_heading(title, level=heading_level)

    content = content.strip()
    if content:
        blocks = re.split(r"\n\s*\n", content)
        for block in blocks:
            block = block.strip()
            if block:
                add_block_as_paragraphs(doc, block)

    if add_page_break:
        doc.add_page_break()


# ============================================================================
# BOQ / COST ANALYSIS
# ============================================================================

def parse_boq_line(line: str) -> Tuple[str, int]:
    """
    Extract description and man-days from BOQ line.
    """
    original = line
    line = line.strip()
    
    line = re.sub(r"^[\-\*•]\s*", "", line)

    man_days = 0
    
    match = re.search(r"(\d+)\s*(?:man[-\s]?days?|days?)", line, flags=re.IGNORECASE)
    if match:
        man_days = int(match.group(1))
    else:
        match = re.search(r"[\(\[](\d+)[\)\]]", line)
        if match:
            man_days = int(match.group(1))
        else:
            match = re.search(r"[→:]\s*(\d+)", line)
            if match:
                man_days = int(match.group(1))
            else:
                match = re.search(r"\s(\d+)\s*$", line)
                if match:
                    man_days = int(match.group(1))

    description = re.split(r"→|:|\(|\[|—|-{2,}|\t", line)[0].strip()
    description = re.sub(r"\s*\d+\s*$", "", description).strip()
    
    if not description:
        description = original.strip()
    
    return description, man_days


def _number_to_indian_words(amount: Decimal) -> str:
    """
    Convert an INR amount to words using Indian numbering system.
    Paise ignored; output ends with 'Rupees Only'.
    """
    n = int(amount.quantize(Decimal("1"), rounding=ROUND_HALF_UP))

    if n == 0:
        return "Zero Rupees Only"

    ones = [
        "", "One", "Two", "Three", "Four", "Five", "Six",
        "Seven", "Eight", "Nine", "Ten", "Eleven", "Twelve",
        "Thirteen", "Fourteen", "Fifteen", "Sixteen",
        "Seventeen", "Eighteen", "Nineteen",
    ]
    tens = [
        "", "", "Twenty", "Thirty", "Forty",
        "Fifty", "Sixty", "Seventy", "Eighty", "Ninety",
    ]

    def two_digit(num: int) -> str:
        if num == 0:
            return ""
        if num < 20:
            return ones[num]
        return (tens[num // 10] + (" " + ones[num % 10] if num % 10 != 0 else "")).strip()

    def three_digit(num: int) -> str:
        if num == 0:
            return ""
        hundred = num // 100
        rest = num % 100
        parts: List[str] = []
        if hundred > 0:
            parts.append(ones[hundred] + " Hundred")
        if rest > 0:
            if parts:
                parts.append("and " + two_digit(rest))
            else:
                parts.append(two_digit(rest))
        return " ".join(parts).strip()

    parts: List[str] = []
    crore = n // 10000000
    n %= 10000000
    if crore:
        parts.append(two_digit(crore) + " Crore")

    lakh = n // 100000
    n %= 100000
    if lakh:
        parts.append(two_digit(lakh) + " Lakh")

    thousand = n // 1000
    n %= 1000
    if thousand:
        parts.append(two_digit(thousand) + " Thousand")

    if n:
        parts.append(three_digit(n))

    words = " ".join(parts).strip()
    return f"{words} Rupees Only"


def add_commercial_boq_section(
    doc: Document,
    title: str,
    content: str,
    rate_usd_per_day: Decimal | int = DEFAULT_RATE_USD,
    exchange_rate: Optional[ExchangeRate] = None,
    heading_level: int = 1,
) -> None:
    """
    Render Commercial Bill of Quantities as a precise cost table.

    Columns:
        Services | Man Days | USD Cost | INR Cost | Total INR Cost

    After the table:
        Subtotal (INR)
        GST @18%
        Final Amount (INR)
        Final Amount in words
    """
    doc.add_heading(title, level=heading_level)

    if exchange_rate is None:
        exchange_rate = fetch_exchange_rate()

    lines = [l for l in content.splitlines() if l.strip()]
    items: List[Tuple[str, int]] = []
    
    default_estimates = {
        "assessment": 5,
        "planning": 5,
        "design": 10,
        "migration": 15,
        "implementation": 15,
        "deployment": 10,
        "configuration": 8,
        "testing": 5,
        "validation": 5,
        "training": 3,
        "documentation": 3,
        "support": 10,
        "optimization": 5,
        "monitoring": 5,
        "license": 0,
        "professional services": 5,
    }
    
    for line in lines:
        if any(keyword in line.lower() for keyword in ["total project cost", "grand total", "subtotal"]):
            continue
            
        description, man_days = parse_boq_line(line)
        
        if man_days == 0 and description:
            desc_lower = description.lower()
            for keyword, default_days in default_estimates.items():
                if keyword in desc_lower:
                    man_days = default_days
                    log.info(
                        f"Assigned default {default_days} man-days to "
                        f"'{description}' based on keyword '{keyword}'"
                    )
                    break
        
        items.append((description, man_days))

    # Table: Services | Man Days | USD Cost | INR Cost | Total INR Cost
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    
    headers = [
        "Services",
        "Man Days",
        "USD Cost",
        "INR Cost",
        "Total INR Cost",
    ]

    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        para = header_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)

    subtotal_inr = Decimal("0.00")
    items_with_cost = 0

    for description, man_days in items:
        row_cells = table.add_row().cells
        
        row_cells[0].text = description
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        if man_days > 0:
            rate_usd = Decimal(rate_usd_per_day)
            desc_lower = description.lower()
            if "workshop" in desc_lower:
                rate_usd = WORKSHOP_RATE_USD

            rate_inr = (rate_usd * exchange_rate.rate).quantize(Decimal("0.01"), ROUND_HALF_UP)
            total_inr = (rate_inr * man_days).quantize(Decimal("0.01"), ROUND_HALF_UP)

            row_cells[1].text = str(man_days)
            row_cells[2].text = format_currency_usd(rate_usd)
            row_cells[3].text = format_currency_inr(rate_inr)
            row_cells[4].text = format_currency_inr(total_inr)

            for i in range(1, 5):
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            subtotal_inr += total_inr
            items_with_cost += 1
        else:
            row_cells[1].text = "N/A"
            row_cells[2].text = "N/A"
            row_cells[3].text = "N/A"
            row_cells[4].text = "N/A"
            
            for i in range(1, 5):
                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.italic = True

    compact_table(table)

    # Subtotal, GST, Final amount
    gst = (subtotal_inr * Decimal("0.18")).quantize(Decimal("0.01"), ROUND_HALF_UP)
    final_amount = (subtotal_inr + gst).quantize(Decimal("0.01"), ROUND_HALF_UP)

    doc.add_paragraph()

    p_sub = doc.add_paragraph()
    r1 = p_sub.add_run("Subtotal (INR): ")
    r1.bold = True
    p_sub.add_run(format_currency_inr(subtotal_inr))

    p_gst = doc.add_paragraph()
    r2 = p_gst.add_run("GST @18% (INR): ")
    r2.bold = True
    p_gst.add_run(format_currency_inr(gst))

    p_total = doc.add_paragraph()
    r3 = p_total.add_run("Final Amount (INR, incl. GST): ")
    r3.bold = True
    r3.font.color.rgb = RGBColor(0, 102, 0)
    p_total.add_run(format_currency_inr(final_amount))

    words = _number_to_indian_words(final_amount)
    p_words = doc.add_paragraph()
    r4 = p_words.add_run("Final Amount in words: ")
    r4.bold = True
    p_words.add_run(words)

    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        f"Note: Exchange rate used: 1 USD = {exchange_rate.rate} INR "
        f"(Source: {exchange_rate.source})."
    )
    note_run.font.size = Pt(9)
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()
    log.info(
        f"BOQ section completed: {items_with_cost}/{len(items)} items costed; "
        f"subtotal={subtotal_inr}, gst={gst}, final={final_amount}"
    )


# ============================================================================
# DETAILED PROPOSAL BUILDER
# ============================================================================

def build_detailed_proposal_docx(sections: Dict[str, str]) -> bytes:
    """
    Build comprehensive detailed proposal document (25-40 pages),
    with grouped sections for ToC (Assessment, Migration, etc.).
    """
    log.info("Building detailed proposal document...")
    
    doc = Document()
    apply_base_styles(doc)

    # Cover page
    cover_text = sections.get("cover_page", "")
    if cover_text:
        customer_name, project_title = add_detailed_cover_page(doc, cover_text, sections)
        log.info(f"Cover page created for {customer_name}: {project_title}")
    else:
        add_branding(doc)
        log.warning("No cover page content provided")

    # Collect section keys (exclude metadata / noise)
    raw_keys = [
        k for k in sections.keys()
        if k not in (
            "cover_page",
            "final_answer",
            "metadata",
            "proposal_type",
            "qa_report",
            "debug_info",
        )
    ]

    # Order keys according to TOC_GROUPS
    ordered_keys: List[str] = []
    seen: set[str] = set()

    for _group_title, group_keys in TOC_GROUPS:
        for k in group_keys:
            if k in raw_keys and k not in seen:
                ordered_keys.append(k)
                seen.add(k)

    # Any leftover keys not mapped to a group
    for k in raw_keys:
        if k not in seen:
            ordered_keys.append(k)
            seen.add(k)

    content_keys = ordered_keys

    # ToC field only (grouping by heading levels)
    add_table_of_contents(doc, [])  # no manual list needed
    log.info(f"Table of Contents placeholder added with {len(content_keys)} sections")

    exchange_rate = fetch_exchange_rate()

    if content_keys:
        last_key = content_keys[-1]
    else:
        last_key = None

    added_groups: set[str] = set()

    for key in content_keys:
        content = sections.get(key, "")
        if not content:
            continue

        group_title = _find_group_for_key(key)
        if group_title and group_title not in added_groups:
            doc.add_heading(group_title, level=1)
            added_groups.add(group_title)

        title = SECTION_TITLES.get(key, key.replace("_", " ").title())
        is_last = (key == last_key)
        heading_level = 2 if group_title else 1

        if key == "commercial_boq_expanded":
            add_commercial_boq_section(
                doc,
                title,
                content,
                rate_usd_per_day=DEFAULT_RATE_USD,
                exchange_rate=exchange_rate,
                heading_level=heading_level,
            )
            log.info(f"BOQ section added: {title}")
        else:
            add_section(
                doc,
                title,
                content,
                add_page_break=not is_last,
                heading_level=heading_level,
            )
            log.info(f"Section added: {title}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    log.info("Detailed proposal document built successfully")
    return buffer.read()


# ============================================================================
# SHORT PROPOSAL BUILDER
# ============================================================================

def extract_price_per_day(service_info: Dict[str, Any]) -> Decimal:
    """
    Extract price per day from service info dictionary.
    """
    candidates = [
        service_info.get("price_man_day"),
        service_info.get("price"),
        service_info.get("price_per_day"),
        service_info.get("cost_per_day"),
        service_info.get("rate"),
    ]
    
    for candidate in candidates:
        if candidate is None:
            continue
        try:
            value = Decimal(str(candidate))
            if value >= 0:
                return value
        except (ValueError, TypeError, ArithmeticError):
            continue

    category = (
        service_info.get("category_name")
        or service_info.get("category")
        or ""
    ).strip()

    service_name = (
        service_info.get("service_name")
        or service_info.get("service")
        or ""
    ).strip()

    if "workshop" in service_name.lower():
        return WORKSHOP_RATE_USD

    if category in CATEGORY_RATE_USD:
        return CATEGORY_RATE_USD[category]

    return DEFAULT_RATE_USD


def build_short_proposal_docx(
    customer_name: str,
    industry: Optional[str] = None,
    deployment_type: Optional[str] = None,
    proposal_type: Optional[str] = "short",
    hardware_choice: Optional[str] = None,
    client_requirements: Optional[str] = None,
    client_boq: Optional[str] = None,
    services: Optional[List[Dict[str, Any]]] = None,
    narrative_sections: Optional[Dict[str, Any]] = None,
) -> str:
    """
    Build concise short-form proposal document (1-3 pages).
    """
    log.info(f"Building short proposal for {customer_name}...")
    
    services = services or []
    narrative_sections = narrative_sections or {}

    exec_summary = narrative_sections.get("executive_summary", "")
    scope_summary = narrative_sections.get("scope_summary", "")
    key_benefits = narrative_sections.get("key_benefits", [])
    risk_note = narrative_sections.get("risk_note", "")
    closing = narrative_sections.get("closing", "")

    doc = Document()
    apply_base_styles(doc)
    add_premium_cover_page(doc, customer_name, industry)

    today_str = datetime.date.today().strftime("%d %B %Y")
    date_para = doc.add_paragraph(f"Date: {today_str}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()  # Spacing

    context_parts: List[str] = []
    context_parts.append(
        f"This {(proposal_type or 'proposal').lower()} is prepared for {customer_name}"
    )
    
    if industry:
        context_parts.append(f"in the {industry} sector")
    if deployment_type:
        context_parts.append(f"with a {deployment_type} deployment model")
    if hardware_choice:
        context_parts.append(f"on {hardware_choice} infrastructure")

    base_context = ", ".join(context_parts) + "."
    
    if client_requirements:
        req_summary = client_requirements.strip().splitlines()[0][:200]
        if req_summary:
            base_context += f' Key objective: "{req_summary.strip()}".'
    
    doc.add_paragraph(base_context)

    if exec_summary:
        doc.add_heading("Executive Summary", level=1)
        for paragraph_text in exec_summary.split("\n\n"):
            text = paragraph_text.strip()
            if text:
                doc.add_paragraph(text)

    if scope_summary:
        doc.add_heading("Scope Summary", level=1)
        for paragraph_text in scope_summary.split("\n\n"):
            text = paragraph_text.strip()
            if text:
                doc.add_paragraph(text)

    if key_benefits:
        if isinstance(key_benefits, str):
            key_benefits = [key_benefits]
        
        doc.add_heading("Key Benefits", level=2)
        for benefit in key_benefits:
            benefit_text = str(benefit).strip()
            if benefit_text:
                doc.add_paragraph(benefit_text, style="List Bullet")

    if risk_note:
        doc.add_heading("Risks & Considerations", level=2)
        for paragraph_text in str(risk_note).split("\n\n"):
            text = paragraph_text.strip()
            if text:
                doc.add_paragraph(text)

    if client_boq:
        doc.add_heading("Client Bill of Quantities (Summary)", level=2)
        for line in client_boq.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Cost Summary", level=1)

    exchange_rate = fetch_exchange_rate()
    
    summary_rows: List[Dict[str, Any]] = []
    grand_total_usd = Decimal("0.00")

    for item in services:
        service_name = (
            item.get("service_name") or 
            item.get("service") or 
            "Service"
        )
        category = (
            item.get("category_name") or 
            item.get("category") or 
            "General Services"
        )
        man_days = int(
            item.get("duration_days") or 
            item.get("duration") or 
            item.get("default_days") or 
            1
        )

        price_per_day = extract_price_per_day(item)
        total_usd = (price_per_day * Decimal(man_days)).quantize(
            Decimal("0.01"), ROUND_HALF_UP
        )

        summary_rows.append({
            "category": category,
            "service": service_name,
            "man_days": man_days,
            "price_per_day_usd": price_per_day,
            "total_usd": total_usd,
        })
        
        grand_total_usd += total_usd

    grand_total_inr = (grand_total_usd * exchange_rate.rate).quantize(
        Decimal("0.01"), ROUND_HALF_UP
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    
    headers = [
        "Category",
        "Service",
        "Man-Days",
        "Rate (USD/day)",
        "Total (USD)",
        "Total (INR)",
    ]

    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        para = header_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header_text)
        run.bold = True

    for row_data in summary_rows:
        row_cells = table.add_row().cells
        
        total_inr = (row_data["total_usd"] * exchange_rate.rate).quantize(
            Decimal("0.01"), ROUND_HALF_UP
        )

        row_cells[0].text = row_data["category"]
        row_cells[1].text = row_data["service"]
        row_cells[2].text = str(row_data["man_days"])
        row_cells[3].text = format_currency_usd(row_data["price_per_day_usd"])
        row_cells[4].text = format_currency_usd(row_data["total_usd"])
        row_cells[5].text = format_currency_inr(total_inr)

    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "GRAND TOTAL"
    total_row[2].text = ""
    total_row[3].text = ""
    total_row[4].text = format_currency_usd(grand_total_usd)
    total_row[5].text = format_currency_inr(grand_total_inr)
    
    for cell in total_row:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.bold = True

    compact_table(table)

    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        f"Exchange rate: 1 USD = {exchange_rate.rate} INR "
        f"({exchange_rate.source})"
    )
    note_run.font.size = Pt(9)
    note_run.italic = True

    doc.add_heading("Terms & Conditions", level=1)
    for term in TERMS_AND_CONDITIONS:
        doc.add_paragraph(term, style="List Bullet")

    doc.add_heading("Closing Note", level=2)
    if closing:
        for paragraph_text in str(closing).split("\n\n"):
            text = paragraph_text.strip()
            if text:
                doc.add_paragraph(text)
    else:
        doc.add_paragraph(
            "We appreciate your consideration and look forward to partnering "
            "with you on this important initiative. Our team is committed to "
            "delivering exceptional results and ensuring your success."
        )
    
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    
    signature_para = doc.add_paragraph("Integrated Tech9Labs Pvt. Ltd. — Professional Services Team")
    for run in signature_para.runs:
        run.bold = True

    filename = f"{sanitize_filename(customer_name)}_Short_Proposal.docx"
    file_path = safe_save_doc(doc, filename)
    
    log.info(f"Short proposal completed: {file_path}")
    return file_path
